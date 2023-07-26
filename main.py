import shutil
import docx.text.font
from github import Github
import docx
from docx.shared import Pt, Inches
from github import Auth
import requests
import pypandoc
import pandoc
from pathlib import Path
from datetime import datetime


repo_name = "InsightSoftwareConsortium/ITK"

# TODO Provide instructions for creating Auth token
# log into github and go to this page https://github.com/settings/tokens
# get token and type into auth_token.txt file
# TODO Read in Auth token from txt file
f = open("my_auth_token.txt")
AUTH_TOKEN = f.readline()[:-1]

# using an access token
auth = Auth.Token(AUTH_TOKEN)

# First create a GitHub instance:

# Public Web GitHub
g = Github(auth=auth)


#  Converts markdown files to docx files
def convert_md_to_docx(file_name: Path, output_file: Path):
    pypandoc.convert_file(
        file_name.as_posix(), to="docx", outputfile=output_file.as_posix()
    )


headers = {
    "Authorization": f"Bearer {AUTH_TOKEN}",
    "X-GitHub-Api-Version": "2022-11-28",
}


# gets comments regarding issue, helper function for get_issues
def get_comments(issue):
    comments_url = issue.comments_url
    response = requests.get(comments_url, headers=headers)
    json_data = response.json()

    comments = []
    for i in range(len(json_data)):
        comment_tuple = (
            str(json_data[i].get("user")["login"]),
            str(json_data[i].get("body")),
        )
        comments.append(comment_tuple)
    return comments


# gets issues from GitHub
def get_issues(repository: str, state: str, updated=None):
    # param rep

    # gets repository
    repo = g.get_repo(repository)

    # stores all issues
    if updated is not None and isinstance(updated, datetime):
        issues = repo.get_issues(state=state, since=updated)
    else:
        issues = repo.get_issues(state=state)
        print(f"{updated} is not a valid type or was not given!")

    # add since parameter feeding into text file
    issues_list = []
    for issue in issues:
        if not issue.pull_request:
            issue_dict = {
                "id_num": issue.number,
                "title": issue.title.replace("`", "'"),
                "body": issue.body,
                "user": issue.user.login,
                "assignees": issue.assignees,
                "published": issue.created_at,
                "labels": [label.name for label in issue.labels if label is not None],
                "milestone": issue.milestone,
                "comments": None,
                "url": issue.url.replace("api.", ""),
                "closed_at": issue.closed_at,
                "updated_at": issue.updated_at,
            }

            if len(issue.assignees) < 1:
                issue_dict["assignees"] = None

            if issue.comments != 0:
                issue_dict["comments"] = get_comments(issue)

            issues_list.append(issue_dict)
    print(len(issues_list))
    return issues_list


#  Creates the markdown files for each issue
def create_md_doc(issue_list, output_dir):
    md_dir = output_dir / "markdown"
    md_dir.mkdir(parents=True, exist_ok=True)

    for issue in issue_list:
        output_name = md_dir / f"issue_{issue['id_num']}.md"

        with open(output_name, "w") as file:
            file.write("\n### Issue:\n")
            file.writelines(
                pypandoc.convert_text(issue["body"], "markdown", format="gfm")
            )

            file.write("\n\n### Comments: \n")
            if issue["comments"]:
                for comment in issue["comments"]:
                    file.write("\n#### " + comment[0] + ":\n")
                    file.writelines(
                        pypandoc.convert_text(comment[1], "markdown", format="gfm")
                    )
            else:
                file.write("\nNo comments at the moment!\n")

            file.write("\n\n### Additional Information: \n")
            file.write("* Milestones: " + str(issue["milestone"]) + "\n")
            file.write("* Assignees: " + str(issue["assignees"]) + "\n")
            url = issue["url"].replace("api.", "")
            file.write("* GitHub Url: " + url.replace("/repos", "") + "\n")
            if issue["closed_at"] is not None:
                file.write("* Closed At: " + str(issue["closed_at"]) + "\n")
            else:
                file.write("* Closed At: Issue Still Open!\n")

            file.write(f"* Last Updated: {issue['updated_at']} ")

            file.write("<br/>\n")
            file.close()


#  Converts folder of Markdown documents into docx files and puts them into new directory
def convert_md_folder(issues, out_dir: Path):
    word_dir = out_dir / "docx"
    md_dir = out_dir / "markdown"
    # state_word_dir = word_dir / state
    word_dir.mkdir(parents=True, exist_ok=True)

    issues_to_convert = [(f'issue_{issue["id_num"]}.md', issue) for issue in issues]

    for tup in issues_to_convert:
        file, issue = tup
        file = list(md_dir.glob(file))[0]
        # file = md_dir / file
        print(file)
        print(issue)
        out_file_name = word_dir / f"{file.stem}.docx"
        convert_md_to_docx(file, out_file_name)
        format_word_doc(file, issue)


#  Adds the heading to every issue's docx file
def format_word_doc(file_name, issue):
    print(file_name, issue["id_num"])
    document = docx.Document(file_name)
    # adds header
    section = document.sections[0]
    heading = section.header
    # heading_para = heading.paragraphs[0]

    table = heading.add_table(rows=2, cols=3, width=Inches(7.5))
    table.alignment = 1
    table.autofit = True
    table.allow_autofit = True
    # table.rows[0].width = Inches(1.0)
    table.columns[2].width = Inches(1.5)
    title_cell = table.rows[0].cells[0]
    title_para = title_cell.paragraphs[0]
    title_para.add_run(
        f"\nIssue No. {issue['id_num']} in The {repo_name[repo_name.index('/') + 1:]} Repository"
    ).bold = True
    title_para.alignment = 1

    issue_cell = table.rows[0].cells[1]
    issue_para = issue_cell.paragraphs[0]
    issue_para.add_run("\n" + issue["title"]).bold = True
    issue_para.alignment = 1

    logo_cell = table.rows[0].cells[2]
    # logo_cell.width = Inches(1.5)
    paragraph = logo_cell.paragraphs[0]
    logo = paragraph.add_run()
    logo.add_picture("BotImageLogo.png", width=Inches(1))  # Image can be changed here
    paragraph.alignment = 1

    date_cell = table.rows[1].cells[0]
    date_para = date_cell.paragraphs[0]
    date_para.add_run("Opened on: ").bold = True
    date_para.add_run(str(issue["published"])).bold = False
    date_para.alignment = 1

    user_cell = table.rows[1].cells[1]
    user_para = user_cell.paragraphs[0]
    user_para.add_run("Opened by: ").bold = True
    user_para.add_run(str(issue["user"])).bold = False
    user_para.alignment = 1

    type_cell = table.rows[1].cells[2]
    # type_cell.width = Inches(1.5)
    type_para = type_cell.paragraphs[0]
    type_para.add_run("Type: ").bold = True
    type_para.add_run(str(issue["labels"])).bold = False
    type_para.alignment = 1

    document.save(file_name)


def validate_state(state_str: str):
    valid_states = ["all", "open", "closed"]
    if state_str.lower() in valid_states:
        return state_str.lower()
    else:
        return valid_states[0]  # defaults to all


#  Function used when gathering issues from a repository for the first time
def initialize_repo(repo_name):
    state = "open"
    state = validate_state(state)

    issues = get_issues(repo_name, state=state)

    out_dir = Path(repo_name + "_issues")
    out_dir.mkdir(parents=True, exist_ok=True)

    create_md_doc(issues[:10], out_dir)
    convert_md_folder(issues, out_dir)

    # for i, file in enumerate(sorted(out_dir.rglob("*.docx"), reverse=True)):
    #     # Now sorts directory with the highest issue numbers first
    #     format_word_doc(file, issues[i])

    f = open(repo_name[repo_name.index("/") + 1 :] + "_time_logs.txt", "a")
    f.write(str(datetime.now()) + "\n")
    f.close()


#  Function used when issues in the repository have already been gathered
def update_repo(repo_name):
    state = "all"
    state = validate_state(state)

    #  Opens time log file and gets the last date this program was run
    with open(repo_name[repo_name.index("/") + 1 :] + "_time_logs.txt", "r") as f:
        for line in f:
            pass
        previous_date = datetime.strptime(line.strip(), "%Y-%m-%d %H:%M:%S.%f")

    #  gets issues from GitHub: state can be "open", "closed" or "all" ... updated parameter will give us the
    #  issues that have been updated since the date passed to the function(usually the last date this program was ran)
    issues = get_issues(repo_name, state=state, updated=previous_date)

    out_dir = Path(repo_name + "_issues")

    create_md_doc(issues, out_dir)
    convert_md_folder(issues, out_dir)

    # for i, issue in enumerate(issues):
    #     # Now sorts directory with the highest issue numbers first
    #     file_name = get_docx_file_name(issues[i], out_dir)
    #     if file_name is not None:
    #         format_word_doc(file_name, issues[i])
    #     else:
    #         print(f"File for issue {issue['id_num']} does not exist")

    with open(repo_name[repo_name.index("/") + 1 :] + "_time_logs.txt", "a") as f:
        f.write(str(datetime.now()) + "\n")


def get_docx_file_name(issue, output_dir):
    issue_num = issue["id_num"]

    file_list = list(output_dir.rglob(f"*{issue_num}.docx"))

    if len(file_list) > 1:
        print(f"More than one file for {issue_num}")
    elif len(file_list) == 0:
        return None
    else:
        return file_list[0]


def clean_up_repo(repo_name):
    repo_dir = Path(repo_name + "_issues")
    num_files = len(list(repo_dir.rglob("*")))
    shutil.rmtree(repo_dir)
    print(f"Removed {num_files} old files from repo directory")


if __name__ == "__main__":
    pypandoc.download_pandoc()

    # clean_up_repo(repo_name)
    #
    # initialize_repo(repo_name)

    update_repo(repo_name)

    print("finished")
